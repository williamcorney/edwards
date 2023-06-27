import docx,os
from Functions import *
from html.parser import HTMLParser
mode = 0

class MyHTMLParser(HTMLParser):
    def handle_starttag(self, tag, attrs):
        global mode
        global paragraph
        global document
        if tag == 'p':     mode = 0
        if tag == 'item':  mode = 0
        if tag == 'h1':    mode = 2
        if tag == 'h3':    mode = 2
        if tag == 'i':     mode = 3
        if tag == 'fnote': mode = 4


    def handle_endtag(self, tag):
        global mode
        global paragraph
        global document
        if tag == 'p':     mode = 1
        if tag == 'item':  mode = 1
        if tag == 'h1':    mode = 1
        if tag == 'h3':    mode = 1
        if tag == 'i':     mode = 1
        if tag == 'fnote': mode = 1


    def handle_data(self, data):
        global mode
        global paragraph
        global document

        if len (data) > 0 :
            if mode == 0: paragraph = document.add_paragraph(data)
            if mode == 1: paragraph.add_run(data).italic = False
            if mode == 2: document.add_heading(data)
            if mode == 3: paragraph.add_run(data).italic = True
            if mode == 4: paragraph.add_footnote(data)

document = docx.Document()
soup = SoupFILEimport('./Volumes/17-01.html')
parser = MyHTMLParser()
parser.feed(str(soup))
os.system("pkill 'Microsoft Word'")
document.save('Volume 17.docx')
