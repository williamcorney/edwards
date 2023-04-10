from bs4 import BeautifulSoup
import os
if os.path.isfile('testing.docx') is True: os.remove('testing.docx')
from tidylib import tidy_document

from html.parser import HTMLParser
from html.entities import name2codepoint
from colorama import Fore, Back, Style
import docx


document = docx.Document()

class MyHTMLParser(HTMLParser):

    def handle_starttag(self, tag, attrs):
        self.fnotecounter = 0
        for attr in attrs:
            self.attribute = (attr[1])
            print(Fore.GREEN + "START (", self.attribute , ") " , tag)
        if tag == 'p':
            self.attribute ='p'


    def handle_endtag(self, tag):
        print (Fore.YELLOW + 'END of ' + tag)
        #print(Fore.RED + "End tag        :",tag[0:10])
    def handle_data(self, data):

        print (Fore.LIGHTBLUE_EX + self.attribute + " " +  data)
        #print(Fore.BLUE + 'Data           : ' + data)

        if 'head' in self.attribute:
            document.add_heading(data, 2)

        if 'fnote' in self.attribute:
            print ('trigger1')

            self.paragraph.add_footnote(data)  # add a footnote
            self.attribute= 'p'
            return
        #else:
           # self.paragraph = document.add_paragraph(data)  # create new paragraph
        if 'p' in self.attribute:
            print ('trigger2')
            self.paragraph = document.add_paragraph(data) # create new paragraph





with open('sermon.html', 'r') as f: contents = f.read()
soup = BeautifulSoup(contents, "lxml")
unwantedtags=['body','html','daterange','date','div']
for tags in unwantedtags:
    target = soup.find_all(tags)
    for items in target: items.unwrap()

for span_tag in soup.findAll('span', {'class':'bibl'}): span_tag.unwrap()
for span_tag in soup.findAll('span', {'class':'reg'}): span_tag.unwrap()
for tag in soup.findAll('span', {'class':''}): tag.unwrap()
for items in soup.select('.fnote'):
    if 'id' in (items.attrs): items.attrs.pop('id')

MyHTMLParser().feed(str(soup))

#print (soup.prettify())


document.save('testing.docx')
