from bs4 import BeautifulSoup
from html.parser import HTMLParser
import docx, os
import subprocess
#pkill - x "Microsoft Word"
subprocess.call('pkill -x "Microsoft Word"', shell=True)

# STAGE 0 ---LOAD UP SOUP FILE
with open('sermon.html', 'r') as f: contents = f.read()
soup = BeautifulSoup(contents, "lxml")
# STAGE 0 ---CLASS SECTION

from colorama import Fore, Back, Style
class MyHTMLParser(HTMLParser):
    def handle_starttag(self, tag, attrs):


        self.count = 0

        if len (attrs) == 0 :
           #print(Fore.LIGHTGREEN_EX + 'start of ' + tag)
           self.flag = tag


        else:
            for item in attrs:

                #print (Fore.LIGHTGREEN_EX + tag + ' (' + item[1] + ')')
                print ('TAG ASSIGNMENT SECTION')
                self.flag = tag + '-' + item[1]
                if item[1] == 'head':
                    print ('TRIGGER ACTIVATED')
                    self.headswitch =1


        #for item in attrs:
            #print ( tag  + item[1] + 'ZZZZZZ' )
        tag = ''

        #for item in attrs:
           # if item[1]

    #def handle_endtag(self, tag):
        #print (Fore.RED + 'end of ' + tag )
    def handle_data(self, data):
        print (Fore.LIGHTGREEN_EX + self.flag)
        print (Fore.LIGHTBLUE_EX + data )

        if self.flag == 'span-head':

            self.paragraph = document.add_heading(data, 3)
            self.paragraph = document.add_paragraph('')

        if self.flag == 'p':
            self.paragraph = document.add_paragraph (data)

        if self.flag == 'span-fnote':

            if self.count == 1:
                self.paragraph.add_run(data)
            if self.count == 0 :
                self.count = self.count + 1
                self.paragraph.add_footnote(data) # add a footnote






            #self.paragraph.add_run(' ')
            #self.paragraph.add_footnote(data) # add a footnote
            #self.paragraph.add_run(' ')
        if self.flag == 'span-font-style:italic':
            if self.count == 1:
                self.paragraph.add_run(data)
            if self.count == 0 :
                self.count = self.count + 1
                self.paragraph.add_run(data).italic=True



        if self.flag == 'b':
            self.paragraph.add_run(data).bold = True
        if self.flag == 'span-bibl':
            if self.count == 1:
                self.paragraph.add_run(data)
            if self.count == 0:
                self.count = self.count + 1
                self.paragraph.add_run(data).italic = True

        #pkill - x "Microsoft Word"


# ******** TRANSFORMATIONS********

# STAGE 1 ---ENTIRELY REMOVE A TAG

decomposetags= ['img','table']
if decomposetags != '':
    for tags in decomposetags:
        NumberofOcurrences = len(soup.find_all(tags))
        if NumberofOcurrences == 1: soup.find(tags).decompose()
        if NumberofOcurrences > 1:
            for items in soup.find_all(tags):
                items.decompose()


# STAGE 2 --- UNWRAP TAGS BY TAG

tagunwrap = ['cit','seg','figure','div']

if len (tagunwrap) == 1 :
    for span_tag in soup.findAll(tagunwrap): span_tag.unwrap()
if len (tagunwrap) > 1 :
    for items in tagunwrap:
        for span_tag in soup.findAll(items): span_tag.unwrap()

# STAGE 3 ---UNWRAP SPAN TAGS BY CLASS

classnames = ['reg','epigraph']
if len (classnames) == 1 :
    for span_tag in soup.findAll('span', {'class': classnames}): span_tag.unwrap()
if len (classnames) > 1 :
    for items in classnames:
        for span_tag in soup.findAll('span', {'class': items}): span_tag.unwrap()


# STAGE 4 - REMOVE ATTRIBUTES FROM CLASS TAG BY NAME OF CLASS TAG AND ATTRIBUTE NAME

thenameofclass ='.fnote'
attrofclass = 'id'

for items in soup.select(thenameofclass):
    if attrofclass in (items.attrs): items.attrs.pop(attrofclass)


#print (soup.prettify())

# STAGE 5 ---CREATE DOCX OBJECT
if os.path.isfile('testing.docx') is True: os.remove('testing.docx')
document = docx.Document()
MyHTMLParser().feed(str(soup))
document.save('testing.docx')


# STAGE 6 - SHOW ALL TAGS IN DOCUMENT
list=[]  #create a list , not for search usage!
taglist = [tag.name for tag in soup.find_all()]
for items in taglist:
    if items not in list:
        list.append(items)
print (list)

#STAGE 7 - SEARCH FOR SPAN TAGS BY  CLASSNAME

searchterm = ['font-style:italic']

searchresponse= (soup.findAll('span', {'class':searchterm}))

print (searchresponse)
print ('hello!!!!')
for items in searchresponse: print (items)

#STAGE 8 - SEARCH FOR SPECIFIC TAG

tagquery=['']
if len(tagquery) < 2:
    print (soup.find_all(tagquery))
if len(tagquery) > 1:
    for items in soup.find_all(tagquery): print (items)
#searchterm = 'span'
#for items in soup.find_all(searchterm): print (items)



subprocess.call('open -a "Microsoft Word" /Users/williamcorney/PycharmProjects/Tags/testing.docx',shell=True)

if os.path.isfile('test.html') is True: os.remove('test.html')
for items in soup:
            with open('test.html', "a") as myfile: myfile.write(str(items))
















"""
# ---UNWRAP TAGS


unwantedtags=['body','html','daterange','date','div','figure']
for tags in unwantedtags:
    target = soup.find_all(tags)
    
    
    for items in target: items.unwrap()

for span_tag in soup.findAll('seg', {'type':'tei wrapper'}): span_tag.unwrap()
for span_tag in soup.findAll('span', {'class':'reg'}): span_tag.unwrap()
for tag in soup.findAll('span', {'class':''}): tag.unwrap()
for items in soup.select('.fnote'):
    if 'id' in (items.attrs): items.attrs.pop('id')



# STAGE 2---REMOVE TAGS - VERSION 2 created accidently

decomposelist= ''
if len (decomposelist) == 1 :
    for tag in decomposelist: soup.find(tag).decompose()
if len (decomposelist) > 1 :
    for tags in decomposelist:
        items = soup.find_all(tags)
        for item in items: item.decompose()



"""

print (soup)
